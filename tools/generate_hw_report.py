"""生成硬件配置方案 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import date

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cell(cell, text, bold=False, size=9, color=None, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_row(row, color_hex):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): color_hex, qn('w:val'): 'clear'})
        shading.append(elm)

def add_table(doc, headers, data, header_color='0F3D3E'):
    t = doc.add_table(rows=1+len(data), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    row0 = t.rows[0]
    shade_row(row0, header_color)
    for i, h in enumerate(headers):
        set_cell(row0.cells[i], h, bold=True, size=9, color=(255,255,255), align='center')
    for i, rd in enumerate(data):
        row = t.rows[i+1]
        if i % 2 == 0:
            shade_row(row, 'F5F5F5')
        for j, val in enumerate(rd):
            set_cell(row.cells[j], str(val), bold=(j==0), size=8)
    return t

# === 标题 ===
title = doc.add_heading('电商内容Agent — 主机配置方案', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'调研日期：{date.today().isoformat()}    版本：v1.0')
doc.add_paragraph('基于项目实际需求和2026年3月市场行情，提供4个可扩展配置方案。所有价格来源于淘宝实际抓取或海外电商公开数据，不确定数据标注"待验证"。')

# PLACEHOLDER_SECTIONS

# === 一、系统需求分析 ===
doc.add_heading('一、系统各环节硬件需求分析', level=1)
doc.add_paragraph('回顾电商内容Agent全流程，逐环节分析是否需要GPU：')

add_table(doc,
    ['环节', '操作', '运行位置', '需要GPU', 'RAM需求'],
    [
        ['竞品采集', 'Playwright爬网页', '本地CPU', '否', '2-4GB'],
        ['抠图(rembg)', 'U2-Net ONNX推理', '本地', '可选', '500MB VRAM'],
        ['场景图生成', '调用Flux/CogView API', '云端', '否', '极低'],
        ['视频生成', '调用Kling/CogVideoX API', '云端', '否', '极低'],
        ['文案生成', '调用LLM API', '云端', '否', '极低'],
        ['详情模块图', 'PIL画布合成', '本地CPU', '否', '低'],
        ['Qdrant向量库', '向量检索', '本地', '否', '1万向量~100MB'],
        ['Redis', '任务队列', '本地', '否', '1-2GB'],
        ['Docker容器', 'Redis+Qdrant+Playwright', '本地', '否', '2-4GB总计'],
        ['Claude Code + Codex', 'AI编程', 'API调用', '否', '500MB/进程'],
        ['FastAPI + React', 'Web服务', '本地CPU', '否', '1-2GB'],
    ])

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('结论：')
r.bold = True
p.add_run('整个系统只有rembg抠图可以用到GPU，其他全部走云端API或纯CPU。GPU是可选项，不是必需品。')

# === 二、rembg CPU vs GPU ===
doc.add_heading('二、rembg 抠图 CPU vs GPU 性能对比', level=1)
doc.add_paragraph('数据来源：ONNX Runtime基准测试(img.ly)、GitHub onnxruntime讨论 #7665')

add_table(doc,
    ['指标', 'CPU (R7 7700X)', 'GPU (RTX 3060 12G)', 'GPU (RTX 5060 Ti 16G)', '说明'],
    [
        ['单张抠图耗时', '2-4秒', '0.2-0.5秒', '<0.1秒', 'GPU快5-20倍'],
        ['批量100张', '3-5分钟', '20-50秒', '<10秒', '批量差距更大'],
        ['VRAM占用', '0(用系统RAM~500MB)', '~500MB-1GB', '~500MB-1GB', '极低'],
        ['首次加载模型', '~3秒', '~5秒(含GPU初始化)', '~3秒', 'GPU首次略慢'],
        ['日处理<50张', '完全够用', '过剩', '过剩', ''],
        ['日处理100-1000张', '偏慢', '舒适', '过剩', ''],
    ])

doc.add_paragraph('')
p2 = doc.add_paragraph()
r2 = p2.add_run('判断：')
r2.bold = True
p2.add_run('日处理50张以内CPU够用。100张以上建议GPU加速。rembg模型只需500MB VRAM，任何入门显卡都够。')

# === 三、DDR5内存市场 ===
doc.add_heading('三、DDR5 内存市场现状（2026年3月）', level=1)
doc.add_paragraph(
    '重要提醒：DDR5内存正处于涨价周期。\n'
    '- Newegg报道：2025年中DDR5价格触底后持续反弹\n'
    '- Overclock3D：2026年DDR5价格持续上涨\n'
    '- Fudzilla：3月出现7.2%月环比下降，部分套装降19%（短期回调）\n'
    '- 分析师预测：2026年内不会大幅降价，最早2027年\n'
    '建议：如果要买内存，趁当前短期回调尽早入手。'
)

# === 四、RTX 5060 Ti ===
doc.add_heading('四、RTX 5060 Ti 已确认信息', level=1)
doc.add_paragraph('来源：Tom\'s Hardware、PC Guide、TechSpot评测、Nvidia官网')

add_table(doc,
    ['参数', 'RTX 5060 Ti 16GB', 'RTX 5060 Ti 8GB', 'RTX 5060'],
    [
        ['MSRP', '$429', '$379', '$299'],
        ['上市日期', '4月16日', '4月16日', '5月'],
        ['VRAM', '16GB GDDR7', '8GB GDDR7', '8GB GDDR7'],
        ['位宽', '128-bit', '128-bit', '128-bit'],
        ['TDP', '180W', '180W', '150W'],
        ['vs 4060 Ti', '+20-22%光栅 +20%光追', '待测', '待测'],
        ['国内预估价', '¥3500-4500(待验证)', '¥3000-3800(待验证)', '¥2500-3000(待验证)'],
    ])

doc.add_paragraph('')
doc.add_paragraph('注意：国内价格为预估，RTX 5060 Ti尚未在国内上市，实际国行定价待验证。')

# PLACEHOLDER_PLANS

# === 五、AM5平台扩展性 ===
doc.add_heading('五、AM5 平台扩展性', level=1)
doc.add_paragraph(
    '来源：AMD官方(Computex 2024)、GameBuilder、Wikipedia\n'
    '- AM5平台AMD官方承诺支持到至少2027年\n'
    '- Zen 5(已发布)和Zen 6(2026-2027发布)都确认兼容AM5\n'
    '- B650M主板PCIe 4.0 x16槽兼容所有RTX 50系列显卡\n'
    '- 部分B650M主板支持最大192GB DDR5(如华硕TUF B650M-E，4 DIMM槽)\n'
    '- 现在买B650M主板，未来CPU和GPU都能升级，不用换主板'
)

# === 六、CPU对比 ===
doc.add_heading('六、CPU 选型对比', level=1)
doc.add_paragraph('来源：UserBenchmark、TechSpot、Tom\'s Hardware')

add_table(doc,
    ['CPU', '核心/线程', 'TDP', '性能差距', '海外价(USD)', '淘宝价(CNY)'],
    [
        ['R7 7700X', '8C/16T', '105W', '基准', '$235-267', '¥575-728'],
        ['R7 9700X', '8C/16T', '65W', '+8%多核,TDP降40%', '~$340', '待验证'],
        ['R5 9600X', '6C/12T', '65W', '接近7700X水平', '~$280', '待验证'],
        ['R7 7800X3D', '8C/16T', '120W', '游戏最强,开发无优势', '$459', '待验证'],
    ])

doc.add_paragraph('')
doc.add_paragraph('建议：R7 7700X性价比最高(淘宝散片¥600)。如果在意功耗和温度，9700X多花¥400-500换来65W TDP。')

# === 七、海外vs国内价格 ===
doc.add_heading('七、海外 vs 国内价格对比', level=1)

add_table(doc,
    ['部件', '海外(Amazon/Newegg)', '国内(淘宝)', '差价结论'],
    [
        ['R7 7700X', '$235-267 (¥1700)', '¥575-728', '国内便宜60%'],
        ['DDR5 32GB 5600', '$88-97 (¥630-700)', '¥653-859', '基本持平'],
        ['DDR5 64GB 5600', '$173 (¥1250)', '¥1538-1694', '国内贵20-30%'],
        ['RTX 3060 12G 全新', '$487 (¥3500)', '¥1329-1379', '国内便宜60%'],
        ['RTX 3060 12G 二手', '$176-275 (¥1270-1980)', '¥490-1160', '国内便宜50%'],
        ['B650M 主板', '$140 (¥1000)', '¥290-709', '国内便宜30-70%'],
        ['1TB NVMe SSD', '$60-80 (¥430-580)', '¥373-459', '基本持平'],
        ['RTX 5060 Ti 16GB', '$429 MSRP (¥3090)', '¥3500-4500(待验证)', '国内可能贵15-30%'],
    ])

doc.add_paragraph('')
doc.add_paragraph('结论：CPU、主板、显卡在国内买明显便宜(散片/淘宝渠道)。内存和SSD中外差不多。新品显卡国内首发可能溢价。')

# === 八、4个配置方案 ===
doc.add_heading('八、4 个配置方案', level=1)

# 方案A
doc.add_heading('方案A：纯CPU版 ~¥2900（先跑起来，随时加卡）', level=2)
add_table(doc,
    ['部件', '选择', '价格', '来源/备注'],
    [
        ['CPU', 'R7 7700X 散片', '¥600', '淘宝实价'],
        ['主板', '技嘉B650M D2HP WiFi(4DIMM,PCIe4.0x16)', '¥499', '淘宝实价'],
        ['内存', 'DDR5 32GB 5600 (16Gx2) 威刚', '¥769', '淘宝实价'],
        ['SSD', '1TB NVMe (金士顿)', '¥373', '淘宝实价'],
        ['电源', '650W铜牌(利民) 预留显卡功耗', '¥203', '淘宝实价'],
        ['机箱', 'M-ATX', '¥150', '淘宝估价'],
        ['散热', '塔式风冷', '¥80', '淘宝估价'],
        ['显卡', '无(用CPU核显)', '¥0', ''],
        ['合计', '', '~¥2674', ''],
    ])

doc.add_paragraph('适合：预算紧张，先把系统跑起来。电源选650W预留显卡余量，主板4DIMM槽可扩展内存。')

# 方案B
doc.add_heading('方案B：入门GPU版 ~¥3700-4200（二手3060，性价比之王）', level=2)
add_table(doc,
    ['部件', '选择', '价格', '来源/备注'],
    [
        ['基础', '同方案A', '¥2674', ''],
        ['显卡', 'RTX 3060 12G 二手', '¥600-1000', '淘宝实价'],
        ['电源', '已含650W,无需更换', '¥0', ''],
        ['合计', '', '~¥3274-3674', ''],
    ])

doc.add_paragraph(
    '优势：12GB VRAM比4060Ti 8G还多，跑rembg/BiRefNet/7B本地模型都够。\n'
    '风险：二手质量不确定，建议选有7天无理由退换的店铺。'
)

# 方案C
doc.add_heading('方案C：主流GPU版 ~¥5500（全新4060Ti）', level=2)
add_table(doc,
    ['部件', '选择', '价格', '来源/备注'],
    [
        ['基础', '同方案A', '¥2674', ''],
        ['显卡', 'RTX 4060 Ti 8G 全新', '¥2599', '淘宝实价'],
        ['合计', '', '~¥5273', ''],
    ])

doc.add_paragraph(
    '优势：全新有保修，功耗低(160W)，DLSS 3支持。\n'
    '劣势：8GB VRAM是硬伤，跑7B本地模型勉强，13B不行。即将被5060Ti替代，性价比不如方案B和D。'
)

# 方案D
doc.add_heading('方案D：性能GPU版 ~¥7500-8500（等5060Ti 16G）', level=2)
add_table(doc,
    ['部件', '选择', '价格', '来源/备注'],
    [
        ['CPU', 'R7 7700X 盒装', '¥728', '淘宝实价'],
        ['主板', '技嘉B650M 小雕WiFi', '¥499', '淘宝实价'],
        ['内存', 'DDR5 64GB 5600 (32Gx2) 威刚', '¥1538', '淘宝实价(¥769x2)'],
        ['SSD', '1TB NVMe', '¥373', '淘宝实价'],
        ['显卡', 'RTX 5060 Ti 16GB', '¥3500-4500', '待验证(MSRP $429)'],
        ['电源', '750W铜牌', '¥250', '淘宝估价'],
        ['机箱', 'M-ATX', '¥200', '淘宝估价'],
        ['散热', '240水冷', '¥250', '淘宝估价'],
        ['合计', '', '~¥7338-8338', ''],
    ])

doc.add_paragraph('适合：预算充足，面向未来。16GB VRAM + 64GB RAM，3-5年不用换。')

# === 九、方案对比 ===
doc.add_heading('九、方案综合对比', level=1)

add_table(doc,
    ['', '方案A 纯CPU', '方案B 二手GPU', '方案C 4060Ti', '方案D 5060Ti'],
    [
        ['总价', '~¥2900', '~¥3700', '~¥5500', '~¥7500-8500'],
        ['VRAM', '0', '12GB', '8GB', '16GB'],
        ['rembg速度', '2-4秒/张', '0.2-0.5秒/张', '0.1-0.3秒/张', '<0.1秒/张'],
        ['本地7B模型', '不行', '可以', '勉强', '舒适'],
        ['本地13B模型', '不行', '不行', '不行', '勉强(量化)'],
        ['BiRefNet抠图', '不行', '可以', '可以', '可以'],
        ['扩展性', '随时加卡', '够用1-2年', 'VRAM瓶颈', '3-5年'],
        ['风险', '无', '二手质量', '即将被替代', '首发可能溢价'],
    ])

# === 十、升级路径 ===
doc.add_heading('十、推荐升级路径', level=1)
doc.add_paragraph(
    '推荐策略：方案A先上 + 4月底加RTX 5060 Ti 16GB\n\n'
    '第一步（现在）：花¥2900买方案A，系统跑起来开始开发\n'
    '第二步（4月底）：RTX 5060 Ti上市价格稳定后，加¥3500-4500显卡\n'
    '第三步（按需）：加¥769内存扩到64GB（Qdrant数据量大时）\n'
    '第四步（2027+）：CPU升级Zen6（主板兼容，不用换）\n\n'
    '总花费：¥2900 + ¥3500-4500 = ¥6400-7400\n'
    '优势：提前一个月开始开发，显卡等首发溢价消退后再买。'
)

# === 数据来源 ===
doc.add_heading('数据来源', level=1)
sources = [
    '淘宝实际抓取价格（2026-03-25，Playwright自动化采集）',
    'Amazon/Newegg 公开价格（2026年3月）',
    'Tom\'s Hardware — RTX 5060 Ti发布报道',
    'PC Guide — RTX 5060/5060 Ti规格确认',
    'TechSpot — RTX 5060 Ti 16GB评测',
    'HyperCyber — RTX 5060 Ti vs 4060 Ti基准测试(+20-22%)',
    'UserBenchmark — R7 7700X vs 9700X对比(+8%)',
    'TechSpot — R7 9700X vs 7700X 45游戏基准测试',
    'Tom\'s Hardware — R5 9600X/R7 9700X评测(TDP 65W)',
    'AMD官方(Computex 2024) — AM5支持到2027年',
    'GameBuilder — Zen 6确认兼容AM5主板',
    'Newegg Insider — DDR5 2026年价格趋势',
    'Overclock3D — DDR5价格持续上涨',
    'Fudzilla — 3月DDR5价格7.2%月环比下降',
    'IQonDigital — RAM价格2026-2028预测',
    'img.ly — ONNX Runtime背景去除基准测试(GPU 20x加速)',
    'Qdrant官方文档 — 容量规划(1万向量~100MB RAM)',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    '声明：所有标注"淘宝实价"的数据来自2026年3月25日自动化抓取。'
    '标注"待验证"的数据为预估值，实际价格以购买时为准。'
    '硬件价格波动频繁，建议购买前再次确认。'
)

# === 保存 ===
output = r'H:\demo2\主机配置方案.docx'
doc.save(output)
print(f'文档已生成: {output}')


