"""
生成 AI API 综合对比 Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import date

doc = Document()

# -- 样式 --
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cell(cell, text, bold=False, size=9, color=None, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_row(row, color_hex):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): color_hex,
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)

def add_table_header(table, headers):
    row = table.rows[0]
    shade_row(row, '0F3D3E')
    for i, h in enumerate(headers):
        set_cell(row.cells[i], h, bold=True, size=9, color=(255,255,255), align='center')

# ============================================================
# 标题
# ============================================================
title = doc.add_heading('电商内容Agent — AI API 综合对比报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'调研日期：{date.today().isoformat()}    版本：v1.0')
doc.add_paragraph(
    '本报告对抠图、生图、视频生成三类 AI API 进行综合对比，'
    '数据来源于各平台官网及开发者文档，价格为调研时点的公开信息。'
)

# ============================================================
# 一、抠图 API
# ============================================================
doc.add_heading('一、抠图（背景去除）API 对比', level=1)
doc.add_paragraph(
    '说明：抠图是将产品从原始背景中分离出来，生成透明/白底图。'
    '电商场景对边缘质量（毛发、透明物体）要求较高。'
)

# PLACEHOLDER_CUTOUT_TABLE

cutout_headers = ['方案', '类型', '质量评级', '价格', 'API文档', '特点', '推荐场景']
cutout_data = [
    ['rembg', '本地开源', '★★★★', '免费（pip install）',
     'github.com/danielgatis/rembg', '离线运行，无限量，U2-Net模型', '默认方案，开发测试'],
    ['remove.bg', '云端API', '★★★★★', '免费50张/月\nAPI: $0.20/张(HD)\n订阅: $9/月起',
     'remove.bg/api', '行业标杆，毛发/透明物体最佳', '高质量商业需求'],
    ['PhotoRoom', '云端API', '★★★★', '$12.99/月起\n批量$0.05/张',
     'photoroom.com/api', '电商专用，批量处理', '电商卖家批量处理'],
    ['Clipdrop\n(Stability AI)', '云端API', '★★★★', '免费额度\nAPI按量计费',
     'clipdrop.co/apis', '多功能AI工具集', '需要多种图片处理'],
    ['BiRefNet', '本地开源', '★★★★★', '免费（需GPU）',
     'github.com/ZhengPeng7/BiRefNet', '边缘处理最精细', '有GPU的高质量场景'],
]

t = doc.add_table(rows=1+len(cutout_data), cols=len(cutout_headers))
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(t, cutout_headers)
for i, row_data in enumerate(cutout_data):
    row = t.rows[i+1]
    if i % 2 == 0:
        shade_row(row, 'F5F5F5')
    for j, val in enumerate(row_data):
        set_cell(row.cells[j], val, bold=(j==0), size=8)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('推荐方案：')
run.bold = True
p.add_run('rembg 本地（默认免费）+ remove.bg API（高质量备选）。已在项目中集成三级降级：rembg → remove.bg → NumPy。')

# ============================================================
# 二、生图 API
# ============================================================
doc.add_heading('二、图像生成 API 对比', level=1)
doc.add_paragraph(
    '说明：用于生成产品场景图（img2img），将白底产品图合成到不同场景中。'
    '关键指标：图生图能力、中文理解、API易用性、价格。'
)

img_headers = ['方案', '提供商', '质量', '价格/张', 'img2img', 'API文档', '国内可用']
img_data = [
    ['Flux.2 Pro', 'fal.ai', '★★★★★',
     '$0.03/MP\n(约$0.03-0.06/张)',
     '支持', 'fal.ai/models/fal-ai/flux-2-pro', '需代理'],
    ['Flux.1 Schnell', 'fal.ai', '★★★★',
     '$0.003/张',
     '支持', 'fal.ai/models/fal-ai/flux/schnell', '需代理'],
    ['Flux.1 Pro', 'fal.ai/官方', '★★★★★',
     '~$0.04/张',
     '支持', 'fal.ai/models/fal-ai/flux-pro', '需代理'],
    ['CogView-4', '智谱AI', '★★★★',
     '~$0.03/张\n(约¥0.2/张)',
     '支持', 'docs.z.ai/guides/image/cogview-4', '直连'],
    ['通义万相 2.6', '阿里云百炼', '★★★★',
     '限时免费体验\n(正式价待定)',
     '支持', 'help.aliyun.com (百炼平台)', '直连'],
    ['DALL-E 3', 'OpenAI', '★★★★',
     '$0.04/张(标准)\n$0.08/张(HD)',
     '不支持', 'platform.openai.com/docs', '需代理'],
    ['GPT-Image-1', 'OpenAI', '★★★★★',
     '$0.02-0.19/张\n(按分辨率)',
     '支持', 'platform.openai.com/docs', '需代理'],
    ['腾讯混元', '腾讯云', '★★★',
     '有免费额度',
     '支持', 'cloud.tencent.cn/document/product/1668', '直连'],
]

t2 = doc.add_table(rows=1+len(img_data), cols=len(img_headers))
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(t2, img_headers)
for i, row_data in enumerate(img_data):
    row = t2.rows[i+1]
    if i % 2 == 0:
        shade_row(row, 'F5F5F5')
    for j, val in enumerate(row_data):
        set_cell(row.cells[j], val, bold=(j==0), size=8)

doc.add_paragraph('')
p2 = doc.add_paragraph()
run2 = p2.add_run('推荐方案：')
run2.bold = True
p2.add_run(
    '性价比最高：Flux Schnell（$0.003/张，通过fal.ai）。'
    '质量最高：Flux.2 Pro（$0.03/张）。'
    '国内直连：智谱 CogView-4（$0.03/张）或通义万相 2.6（限免）。'
)

# PLACEHOLDER_VIDEO_TABLE

# ============================================================
# 三、视频生成 API
# ============================================================
doc.add_heading('三、视频生成 API 对比', level=1)
doc.add_paragraph(
    '说明：用于将产品场景图/描述生成短视频。'
    '关键指标：视频质量、时长、物理真实感、价格、API可用性。'
)

vid_headers = ['方案', '提供商', '质量', '价格', '时长', 'API文档', '国内可用']
vid_data = [
    ['Kling 2.6 Pro', '快手', '★★★★★',
     '$0.07/秒(无音频)\n$0.14/秒(含音频)\n(5秒≈$0.35-0.70)',
     '5-10秒', 'fal.ai/models/fal-ai/kling-video', '直连/fal.ai'],
    ['Kling 3.0', '快手', '★★★★★',
     '第三方约$0.075/秒\n(官方$4200/月起)',
     '5-10秒', 'kling3api.com / klingai.com', '直连'],
    ['Runway Gen-4.5', 'Runway', '★★★★★',
     '订阅$12-76/月\nAPI按秒计费',
     '5-10秒', 'docs.dev.runwayml.com/api', '需代理'],
    ['MiniMax\nHailuo 02', 'MiniMax', '★★★★',
     '720p 25fps\n约$0.3-0.5/视频',
     '5-6秒', 'platform.minimax.io', '直连'],
    ['CogVideoX-3', '智谱AI', '★★★★',
     '约$0.32/视频\n(Replicate)',
     '6秒', 'docs.z.ai/guides/video/cogvideox-3', '直连'],
    ['Sora 2', 'OpenAI', '★★★★★',
     '$6-15/视频',
     '5-20秒', 'platform.openai.com', '需代理'],
    ['Veo 3.1', 'Google', '★★★★★',
     '$3-15/视频\nAI Studio免费',
     '5-8秒', 'ai.google.dev', '需代理'],
    ['通义万相\nWan 2.6', '阿里云', '★★★★',
     '限时免费体验',
     '5-10秒', 'help.aliyun.com (百炼)', '直连'],
]

t3 = doc.add_table(rows=1+len(vid_data), cols=len(vid_headers))
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(t3, vid_headers)
for i, row_data in enumerate(vid_data):
    row = t3.rows[i+1]
    if i % 2 == 0:
        shade_row(row, 'F5F5F5')
    for j, val in enumerate(row_data):
        set_cell(row.cells[j], val, bold=(j==0), size=8)

doc.add_paragraph('')
p3 = doc.add_paragraph()
run3 = p3.add_run('推荐方案：')
run3.bold = True
p3.add_run(
    '性价比最高：Kling 2.6（通过fal.ai，$0.07/秒）。'
    '国内直连：智谱 CogVideoX-3 或 MiniMax Hailuo。'
    '质量最高：Sora 2（但$6-15/视频，成本高）。'
)

# ============================================================
# 四、免费试用额度汇总
# ============================================================
doc.add_heading('四、免费试用 / 免费额度汇总', level=1)
doc.add_paragraph('以下为各平台提供的免费试用额度，适合前期测试和验证效果。')

free_headers = ['平台', '类型', '免费额度', '是否需信用卡', '备注']
free_data = [
    ['rembg', '抠图', '完全免费，无限量', '否', '本地运行，pip install'],
    ['remove.bg', '抠图', '50张/月（预览质量）', '否', 'API注册即可'],
    ['fal.ai', '生图/视频', '注册送$10免费额度', '否', '可用于Flux/Kling/MiniMax等所有模型'],
    ['智谱AI', '生图/视频', '注册送免费token', '否', 'CogView-4 + CogVideoX'],
    ['阿里云百炼', '生图/视频', '通义万相限时免费体验', '否', '需阿里云账号'],
    ['腾讯混元', '生图', '有免费额度', '否', '需腾讯云账号'],
    ['OpenAI', '生图', '无免费API额度（2026年）', '是', '需预充值'],
    ['Runway', '视频', '注册送125积分', '否', '约可生成几个短视频'],
    ['Kling AI', '视频', '免费版每日有限额度', '否', '标准版$6.99/月起'],
    ['Google AI Studio', '视频', 'Veo 3.1 免费使用', '否', '需Google账号，国内需代理'],
    ['MiniMax', '视频', 'Hailuo免费版3个/天', '否', 'Pro版$9.99/月'],
]

t4 = doc.add_table(rows=1+len(free_data), cols=len(free_headers))
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(t4, free_headers)
for i, row_data in enumerate(free_data):
    row = t4.rows[i+1]
    if i % 2 == 0:
        shade_row(row, 'F5F5F5')
    for j, val in enumerate(row_data):
        set_cell(row.cells[j], val, bold=(j==0), size=8)

# ============================================================
# 五、综合推荐方案
# ============================================================
doc.add_heading('五、综合推荐方案', level=1)
doc.add_paragraph('基于项目需求（面向国内电商、需要国内直连、控制成本），推荐以下组合：')

rec_headers = ['功能', '首选方案', '首选价格', '备选方案', '备选价格', '选择理由']
rec_data = [
    ['抠图', 'rembg（本地）', '免费', 'remove.bg API', '$0.20/张',
     '默认免费无限量，高质量需求切换remove.bg'],
    ['生图\n（场景图）', 'Flux.2 Pro\n(fal.ai)', '$0.03/张', '智谱 CogView-4', '~$0.03/张',
     'Flux质量最高；CogView国内直连'],
    ['生图\n（批量/快速）', 'Flux Schnell\n(fal.ai)', '$0.003/张', '通义万相 2.6', '限免',
     '极低成本批量生成；万相国内直连'],
    ['视频生成', 'Kling 2.6\n(fal.ai)', '$0.07/秒', '智谱 CogVideoX-3', '~$0.32/视频',
     'Kling性价比+质量最优；CogVideoX国内直连'],
]

t5 = doc.add_table(rows=1+len(rec_data), cols=len(rec_headers))
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(t5, rec_headers)
for i, row_data in enumerate(rec_data):
    row = t5.rows[i+1]
    if i % 2 == 0:
        shade_row(row, 'F5F5F5')
    for j, val in enumerate(row_data):
        set_cell(row.cells[j], val, bold=(j==0), size=8)

# ============================================================
# 六、成本估算
# ============================================================
doc.add_heading('六、月度成本估算（以1000个产品/月为例）', level=1)

cost_headers = ['环节', '数量', '单价', '月成本', '方案']
cost_data = [
    ['抠图', '1000张', '免费', '$0', 'rembg本地'],
    ['场景图生成', '4000张\n(每产品4张)', '$0.003/张', '$12', 'Flux Schnell'],
    ['场景图生成\n（高质量）', '4000张', '$0.03/张', '$120', 'Flux.2 Pro'],
    ['视频生成', '1000个\n(5秒/个)', '$0.35/个', '$350', 'Kling 2.6 (fal.ai)'],
    ['合计（经济版）', '', '', '~$362/月', 'Schnell + Kling'],
    ['合计（高质量版）', '', '', '~$470/月', 'Flux Pro + Kling'],
]

t6 = doc.add_table(rows=1+len(cost_data), cols=len(cost_headers))
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(t6, cost_headers)
for i, row_data in enumerate(cost_data):
    row = t6.rows[i+1]
    if i % 2 == 0:
        shade_row(row, 'F5F5F5')
    for j, val in enumerate(row_data):
        is_total = '合计' in val if j == 0 else False
        set_cell(row.cells[j], val, bold=(j==0 or is_total), size=8)

# ============================================================
# 数据来源
# ============================================================
doc.add_heading('数据来源', level=1)
sources = [
    'fal.ai 官方定价页 — fal.ai/pricing',
    'remove.bg API 页面 — remove.bg/api（免费50张/月，HD $0.20/张）',
    'OpenAI API 定价 — openai.com/api/pricing（DALL-E 3: $0.04-0.12/张）',
    'Kling AI via fal.ai — fal.ai/models/fal-ai/kling-video（$0.07/秒 Pro）',
    '智谱AI开放平台 — docs.z.ai（CogView-4: ~$0.03/张）',
    '阿里云百炼 — help.aliyun.com/zh/model-studio（通义万相限免）',
    'Runway API 文档 — docs.dev.runwayml.com/api',
    'MiniMax 平台 — platform.minimax.io',
    'Flux 模型对比 — techlifeadventures.com（Flux Pro ~$0.04/张, Schnell ~$0.003/张）',
    'AI API 综合对比 — teamday.ai/blog/ai-api-pricing-comparison-2026',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    '注意：以上价格为调研时点（2026年3月）的公开信息，实际价格可能随平台调整而变化。'
    '建议在接入前到各平台官网确认最新定价。',
)

# ============================================================
# 保存
# ============================================================
output_path = r'H:\demo2\AI_API综合对比报告.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
