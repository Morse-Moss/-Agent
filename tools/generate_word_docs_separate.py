#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate separate Word documents from markdown files"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def add_markdown_to_doc(doc, markdown_text):
    """Convert markdown text to Word document elements"""

    lines = markdown_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Handle headers
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='Heading 1')
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(6)
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:], style='Heading 2')
            p_format = p.paragraph_format
            p_format.space_before = Pt(10)
            p_format.space_after = Pt(4)
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 3')
            p_format = p.paragraph_format
            p_format.space_before = Pt(8)
            p_format.space_after = Pt(2)
        elif line.startswith('#### '):
            p = doc.add_paragraph(line[5:], style='Heading 4')

        # Handle tables
        elif line.startswith('| '):
            # Collect table rows
            table_rows = []
            while i < len(lines) and lines[i].startswith('| '):
                row = lines[i].split('|')[1:-1]  # Remove first and last empty elements
                row = [cell.strip() for cell in row]
                table_rows.append(row)
                i += 1

            if table_rows:
                # Skip separator row if present
                if len(table_rows) > 1 and all(c.strip().replace('-', '').replace(' ', '') == '' for c in table_rows[1]):
                    table_rows = [table_rows[0]] + table_rows[2:]

                # Create table
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        if row_idx == 0:
                            # Header row formatting
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            i -= 1  # Adjust because we'll increment at the end

        # Handle bullet points
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')

        # Handle numbered lists
        elif line and line[0].isdigit() and '. ' in line:
            parts = line.split('. ', 1)
            if len(parts) == 2:
                p = doc.add_paragraph(parts[1], style='List Number')

        # Regular paragraph
        else:
            if line.strip():
                p = doc.add_paragraph(line)

        i += 1

def create_document(file_path, output_path, title):
    """Create a Word document from a markdown file"""

    if not os.path.exists(file_path):
        print(f'File not found: {file_path}')
        return

    # Create document
    doc = Document()

    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    date_run.font.size = Pt(11)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Read and add markdown content
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    add_markdown_to_doc(doc, content)

    # Save document
    doc.save(output_path)
    print(f'Generated: {output_path}')

def main():
    files = [
        ('g:\\demo2\\docs\\internal\\source\\产品需求草案_v0.5.md', 'g:\\demo2\\产品需求草案_v0.5.docx', '产品需求草案 v0.5'),
        ('g:\\demo2\\docs\\internal\\source\\技术实现方案_v0.5.md', 'g:\\demo2\\技术实现方案_v0.5.docx', '技术实现方案 v0.5'),
        ('g:\\demo2\\docs\\internal\\source\\可复用分析_电商内容Agent_v0.5.md', 'g:\\demo2\\可复用分析_电商内容Agent_v0.5.docx', '可复用分析：电商内容Agent v0.5'),
    ]

    for input_file, output_file, title in files:
        create_document(input_file, output_file, title)

if __name__ == '__main__':
    main()
