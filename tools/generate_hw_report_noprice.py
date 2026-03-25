"""生成硬件配置方案 Word 文档 — 无价格版（可公开分享）"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import date

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cell(cell, text, bold=False, size=9, color=None, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_row(row, color_hex):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): color_hex, qn('w:val'): 'clear'})
        shading.append(elm)

def add_table(doc, headers, data, header_color='0F3D3E'):
    t = doc.add_table(rows=1+len(data), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    row0 = t.rows[0]
    shade_row(row0, header_color)
    for i, h in enumerate(headers):
        set_cell(row0.cells[i], h, bold=True, size=9, color=(255,255,255), align='center')
    for i, rd in enumerate(data):
        row = t.rows[i+1]
        if i % 2 == 0:
            shade_row(row, 'F5F5F5')
        for j, val in enumerate(rd):
            set_cell(row.cells[j], str(val), bold=(j==0), size=8)
    return t

# === 标题 ===
title = doc.add_heading('电商内容Agent — 主机配置方案（技术选型版）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'调研日期：{date.today().isoformat()}    版本：v1.0')
doc.add_paragraph('基于项目实际需求，提供4个可扩展配置方案。本文档聚焦技术选型和性能分析，不含具体价格。')

# === 一、系统需求分析 ===
doc.add_heading('一、系统各环节硬件需求分析', level=1)
doc.add_paragraph('回顾电商内容Agent全流程，逐环节分析是否需要GPU：')

add_table(doc,
    ['环节', '操作', '运行位置', '需要GPU', 'RAM需求'],
    [
        ['竞品采集', 'Playwright爬网页', '本地CPU', '否', '2-4GB'],
        ['抠图(rembg)', 'U2-Net ONNX推理', '本地', '可选', '500MB VRAM'],
        ['场景图生成', '调用Flux/CogView API', '云端', '否', '极低'],
        ['视频生成', '调用Kling/CogVideoX API', '云端', '否', '极低'],
        ['文案生成', '调用LLM API', '云端', '否', '极低'],
        ['详情模块图', 'PIL画布合成', '本地CPU', '否', '低'],
        ['Qdrant向量库', '向量检索', '本地', '否', '1万向量~100MB'],
        ['Redis', '任务队列', '本地', '否', '1-2GB'],
        ['Docker容器', 'Redis+Qdrant+Playwright', '本地', '否', '2-4GB总计'],
        ['Claude Code + Codex', 'AI编程', 'API调用', '否', '500MB/进程'],
        ['FastAPI + React', 'Web服务', '本地CPU', '否', '1-2GB'],
    ])

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('结论：')
r.bold = True
p.add_run('整个系统只有rembg抠图可以用到GPU，其他全部走云端API或纯CPU。GPU是可选项，不是必需品。')

# === 二、rembg CPU vs GPU ===
doc.add_heading('二、rembg 抠图 CPU vs GPU 性能对比', level=1)
doc.add_paragraph('数据来源：ONNX Runtime基准测试(img.ly)、GitHub onnxruntime讨论 #7665')

add_table(doc,
    ['指标', 'CPU (8核16线程)', 'GPU (12GB VRAM)', 'GPU (16GB VRAM)', '说明'],
    [
        ['单张抠图耗时', '2-4秒', '0.2-0.5秒', '<0.1秒', 'GPU快5-20倍'],
        ['批量100张', '3-5分钟', '20-50秒', '<10秒', '批量差距更大'],
        ['VRAM占用', '0(用系统RAM~500MB)', '~500MB-1GB', '~500MB-1GB', '极低'],
        ['首次加载模型', '~3秒', '~5秒(含GPU初始化)', '~3秒', 'GPU首次略慢'],
        ['日处理<50张', '完全够用', '过剩', '过剩', ''],
        ['日处理100-1000张', '偏慢', '舒适', '过剩', ''],
    ])

doc.add_paragraph('')
p2 = doc.add_paragraph()
r2 = p2.add_run('判断：')
r2.bold = True
p2.add_run('日处理50张以内CPU够用。100张以上建议GPU加速。rembg模型只需500MB VRAM，任何入门显卡都够。')

# === 三、DDR5内存市场 ===
doc.add_heading('三、DDR5 内存市场现状（2026年3月）', level=1)
doc.add_paragraph(
    '重要提醒：DDR5内存正处于涨价周期。\n'
    '- Newegg报道：2025年中DDR5价格触底后持续反弹\n'
    '- Overclock3D：2026年DDR5价格持续上涨\n'
    '- Fudzilla：3月出现7.2%月环比下降，部分套装降19%（短期回调）\n'
    '- 分析师预测：2026年内不会大幅降价，最早2027年\n'
    '建议：如果要买内存，趁当前短期回调尽早入手。'
)

# === 四、RTX 5060 Ti ===
doc.add_heading('四、RTX 5060 Ti 已确认规格', level=1)
doc.add_paragraph('来源：Tom\'s Hardware、PC Guide、TechSpot评测、Nvidia官网')

add_table(doc,
    ['参数', 'RTX 5060 Ti 16GB', 'RTX 5060 Ti 8GB', 'RTX 5060'],
    [
        ['上市日期', '4月16日', '4月16日', '5月'],
        ['VRAM', '16GB GDDR7', '8GB GDDR7', '8GB GDDR7'],
        ['位宽', '128-bit', '128-bit', '128-bit'],
        ['TDP', '180W', '180W', '150W'],
        ['架构', 'Blackwell GB206', 'Blackwell GB206', 'Blackwell GB206'],
        ['vs 4060 Ti', '+20-22%光栅 +20%光追', '待测', '待测'],
        ['DLSS', 'DLSS 4 (多帧生成)', 'DLSS 4', 'DLSS 4'],
    ])

# === 五、AM5平台扩展性 ===
doc.add_heading('五、AM5 平台扩展性', level=1)
doc.add_paragraph(
    '来源：AMD官方(Computex 2024)、GameBuilder、Wikipedia\n'
    '- AM5平台AMD官方承诺支持到至少2027年\n'
    '- Zen 5(已发布)和Zen 6(2026-2027发布)都确认兼容AM5\n'
    '- B650M主板PCIe 4.0 x16槽兼容所有RTX 50系列显卡\n'
    '- 部分B650M主板支持最大192GB DDR5(如华硕TUF B650M-E，4 DIMM槽)\n'
    '- 现在买B650M主板，未来CPU和GPU都能升级，不用换主板'
)

# === 六、CPU对比 ===
doc.add_heading('六、CPU 选型对比', level=1)
doc.add_paragraph('来源：UserBenchmark、TechSpot、Tom\'s Hardware')

add_table(doc,
    ['CPU', '核心/线程', 'TDP', '架构', '性能差距', '适合场景'],
    [
        ['R7 7700X', '8C/16T', '105W', 'Zen 4', '基准', '性价比首选'],
        ['R7 9700X', '8C/16T', '65W', 'Zen 5', '+8%多核,TDP降40%', '低功耗静音'],
        ['R5 9600X', '6C/12T', '65W', 'Zen 5', '接近7700X水平', '预算有限'],
        ['R7 7800X3D', '8C/16T', '120W', 'Zen 4 + 3D V-Cache', '游戏最强', '游戏为主(开发无优势)'],
    ])

doc.add_paragraph('')
doc.add_paragraph('建议：R7 7700X性价比最高。如果在意功耗和温度，9700X用65W TDP达到接近性能。')

# === 七、显卡选型对比 ===
doc.add_heading('七、显卡选型对比', level=1)

add_table(doc,
    ['显卡', 'VRAM', 'TDP', '架构', 'rembg加速', '本地7B模型', '本地13B模型', '扩展性'],
    [
        ['无(核显)', '0', '0W', '-', '不支持', '不行', '不行', '随时加卡'],
        ['RTX 3060 12G', '12GB GDDR6', '170W', 'Ampere', '5-10x', '可以', '不行', '够用1-2年'],
        ['RTX 4060 Ti 8G', '8GB GDDR6', '160W', 'Ada', '10-15x', '勉强', '不行', 'VRAM瓶颈'],
        ['RTX 5060 Ti 16G', '16GB GDDR7', '180W', 'Blackwell', '15-20x', '舒适', '勉强(量化)', '3-5年'],
        ['RTX 5060 8G', '8GB GDDR7', '150W', 'Blackwell', '10-15x', '勉强', '不行', 'VRAM瓶颈'],
    ])

doc.add_paragraph('')
doc.add_paragraph('关键：VRAM容量决定能跑什么模型。12GB是实用门槛，16GB是舒适线。8GB在2026年已经偏小。')

# === 八、4个配置方案 ===
doc.add_heading('八、4 个配置方案', level=1)

# 方案A
doc.add_heading('方案A：纯CPU版（先跑起来，随时加卡）', level=2)
add_table(doc,
    ['部件', '推荐选择', '备注'],
    [
        ['CPU', 'AMD R7 7700X', '8核16线程，AM5平台'],
        ['主板', 'B650M (4 DIMM槽，PCIe 4.0 x16)', '如技嘉B650M D2HP WiFi'],
        ['内存', 'DDR5 32GB 5600 (16Gx2)', '留2槽可扩展到64/128GB'],
        ['SSD', '1TB NVMe PCIe 4.0', '主板有2个M.2槽可加第二块'],
        ['电源', '650W 铜牌', '预留显卡功耗余量(5060Ti TDP 180W)'],
        ['显卡', '无(核显)', 'rembg用CPU跑，2-4秒/张'],
        ['散热', '塔式风冷', ''],
    ])
doc.add_paragraph('适合：预算有限，先把系统跑起来。电源和主板预留扩展空间。')

# 方案B
doc.add_heading('方案B：入门GPU版（二手3060，性价比之王）', level=2)
add_table(doc,
    ['部件', '推荐选择', '备注'],
    [
        ['基础', '同方案A', ''],
        ['显卡', 'RTX 3060 12G (二手)', '12GB VRAM，可跑BiRefNet和7B本地模型'],
    ])
doc.add_paragraph(
    '优势：12GB VRAM比4060Ti 8G还多，覆盖所有当前需求。\n'
    '风险：二手质量不确定，建议选有退换保障的店铺。'
)

# 方案C
doc.add_heading('方案C：主流GPU版（全新4060Ti）', level=2)
add_table(doc,
    ['部件', '推荐选择', '备注'],
    [
        ['基础', '同方案A', ''],
        ['显卡', 'RTX 4060 Ti 8G (全新)', '全新有保修，功耗低，DLSS 3'],
    ])
doc.add_paragraph(
    '优势：全新保修，功耗低(160W)。\n'
    '劣势：8GB VRAM是硬伤，即将被5060Ti替代。不推荐，除非急需且不想等。'
)

# 方案D
doc.add_heading('方案D：性能GPU版（等5060Ti 16G，推荐）', level=2)
add_table(doc,
    ['部件', '推荐选择', '备注'],
    [
        ['CPU', 'AMD R7 7700X', '盒装或散片均可'],
        ['主板', 'B650M (4 DIMM槽)', '如技嘉B650M 小雕WiFi'],
        ['内存', 'DDR5 64GB 5600 (32Gx2)', 'Qdrant+Docker+多Agent舒适'],
        ['SSD', '1TB NVMe PCIe 4.0', ''],
        ['显卡', 'RTX 5060 Ti 16GB (4月16日上市)', '16GB VRAM，Blackwell架构'],
        ['电源', '750W 铜牌', '预留升级余量'],
        ['散热', '240水冷', ''],
    ])
doc.add_paragraph('适合：预算充足，面向未来。16GB VRAM + 64GB RAM，3-5年不用换。')

# === 九、方案对比 ===
doc.add_heading('九、方案综合对比', level=1)

add_table(doc,
    ['', '方案A 纯CPU', '方案B 二手GPU', '方案C 4060Ti', '方案D 5060Ti'],
    [
        ['VRAM', '0', '12GB', '8GB', '16GB'],
        ['系统RAM', '32GB', '32GB', '32GB', '64GB'],
        ['rembg速度', '2-4秒/张', '0.2-0.5秒/张', '0.1-0.3秒/张', '<0.1秒/张'],
        ['本地7B模型', '不行', '可以', '勉强', '舒适'],
        ['本地13B模型', '不行', '不行', '不行', '勉强(量化)'],
        ['BiRefNet抠图', '不行', '可以', '可以', '可以'],
        ['Qdrant 10万向量', '够用', '够用', '够用', '舒适(64GB RAM)'],
        ['多Agent并行', '够用', '够用', '够用', '舒适(64GB RAM)'],
        ['扩展性', '随时加卡', '够用1-2年', 'VRAM瓶颈', '3-5年'],
        ['风险', '无', '二手质量', '即将被替代', '首发可能溢价'],
    ])

# === 十、升级路径 ===
doc.add_heading('十、推荐升级路径', level=1)
doc.add_paragraph(
    '推荐策略：方案A先上 + 等RTX 5060 Ti 16GB上市后升级\n\n'
    '第一步（现在）：买方案A，系统跑起来开始开发\n'
    '第二步（4月底）：RTX 5060 Ti上市价格稳定后加显卡\n'
    '第三步（按需）：加内存扩到64GB（Qdrant数据量大时）\n'
    '第四步（2027+）：CPU升级Zen6（主板兼容，不用换）\n\n'
    '优势：提前一个月开始开发，显卡等首发溢价消退后再买。\n'
    'AM5平台支持到2027年+，主板一次投资长期受益。'
)

# === 数据来源 ===
doc.add_heading('数据来源', level=1)
sources = [
    'Tom\'s Hardware — RTX 5060 Ti发布报道及规格确认',
    'PC Guide — RTX 5060/5060 Ti上市日期和配置',
    'TechSpot — RTX 5060 Ti 16GB评测',
    'HyperCyber — RTX 5060 Ti vs 4060 Ti基准测试(+20-22%)',
    'UserBenchmark — R7 7700X vs 9700X对比(+8%)',
    'TechSpot — R7 9700X vs 7700X 45游戏基准测试',
    'Tom\'s Hardware — R5 9600X/R7 9700X评测(TDP 65W)',
    'AMD官方(Computex 2024) — AM5支持到2027年',
    'GameBuilder — Zen 6确认兼容AM5主板',
    'Newegg Insider — DDR5 2026年价格趋势',
    'Overclock3D — DDR5价格持续上涨',
    'Fudzilla — 3月DDR5价格7.2%月环比下降',
    'IQonDigital — RAM价格2026-2028预测',
    'img.ly — ONNX Runtime背景去除基准测试(GPU 20x加速)',
    'Qdrant官方文档 — 容量规划(1万向量~100MB RAM)',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('声明：本文档聚焦技术选型分析，不含具体价格。硬件价格波动频繁，请购买前自行查询最新行情。')

# === 保存 ===
output = r'H:\demo2\主机配置方案_技术选型版.docx'
doc.save(output)
print(f'文档已生成: {output}')
